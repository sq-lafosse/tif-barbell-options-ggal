"""generate_strategy_memo.py — Genera el memo Word con la descripción completa de la
Estrategia Barbell, sus supuestos, las métricas del backtest y la comparación vs.
benchmarks (ADR GGAL, Merval USD).

No es parte del pipeline de datos (no lee `data/raw/`); solo lee los CSV ya procesados
por `backtest.py`/`benchmark.py`/`report.py` y `config.yaml`, y redacta el documento.
Pensado como insumo de lectura para el tutor / capítulo de resultados de la tesis, no
como artefacto que se regenera automáticamente en cada corrida del pipeline.
"""

from pathlib import Path

import pandas as pd
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
CONFIG_PATH = ROOT / "config.yaml"
FULL_COMPARISON_PATH = ROOT / "data/processed/full_comparison.csv"
BACKTEST_COMPARISON_PATH = ROOT / "data/processed/backtest_comparison.csv"
OUTPUT_PATH = ROOT / "reports" / "Estrategia_Barbell_GGAL_Memo.docx"

ACCENT = RGBColor(0x1B, 0x4F, 0x72)


def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _table_from_df(doc: Document, df: pd.DataFrame, pct_cols: list[str] | None = None) -> None:
    pct_cols = pct_cols or []
    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for j, col in enumerate(df.columns):
        hdr[j + 1].text = col
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx)
        for j, col in enumerate(df.columns):
            val = row[col]
            if pd.isna(val):
                cells[j + 1].text = "—"
            elif col in pct_cols:
                cells[j + 1].text = f"{val:,.2f}%"
            elif isinstance(val, float):
                cells[j + 1].text = f"{val:,.2f}"
            else:
                cells[j + 1].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)


def main() -> int:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    full = pd.read_csv(FULL_COMPARISON_PATH, index_col="model")
    bt = pd.read_csv(BACKTEST_COMPARISON_PATH, index_col="model")

    hybrid_full = full.loc["Barbell (hybrid)"]
    adr_full = full.loc["GGAL ADR buy-and-hold"]
    merval_full = full.loc["Merval en USD"]
    hybrid_bt = bt.loc["hybrid"]

    doc = Document()

    # Estilo de fuente base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Estrategia Barbell sobre Opciones de GGAL", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Descripción de la estrategia, supuestos metodológicos, métricas del backtest "
        "y comparación contra benchmarks (ADR GGAL y Merval en USD)"
    )
    sub_run.italic = True

    # -------------------------------------------------------------------
    # 1. Qué es la estrategia
    # -------------------------------------------------------------------
    _heading(doc, "1. Descripción de la estrategia", 1)
    _para(
        doc,
        "La Barbell es una asignación bimodal del capital, pensada para evitar el "
        "“medio frágil”: en vez de invertir todo en una posición de riesgo medio "
        "(buy-and-hold lineal del activo), el capital se divide en dos polos con perfiles "
        "de riesgo extremos y complementarios.",
    )

    pf = config["portfolio"]
    moneyness = config["moneyness"]

    _heading(doc, "1.1 Polo seguro", 2)
    _bullet(
        doc,
        f"Peso objetivo: {pf['safe_weight']*100:.0f}% del capital total (parámetro "
        "tentativo, sujeto a sensibilización — ver sección 2).",
    )
    _bullet(
        doc,
        "Instrumento: T-Bills 3M (tasa libre de riesgo, fuente FRED), devengando interés "
        "día a día sobre el capital no invertido en el polo agresivo.",
    )
    _bullet(
        doc,
        "Función: inmunizar el capital ante eventos de cola izquierda — el polo seguro "
        "nunca se expone a la caída del activo de riesgo.",
    )

    _heading(doc, "1.2 Polo agresivo", 2)
    _bullet(
        doc,
        f"Peso objetivo: {pf['aggressive_weight']*100:.0f}% del capital total (tentativo).",
    )
    _bullet(
        doc,
        "Instrumento: compra sistemática de opciones Put Out-of-the-Money (OTM) sobre "
        "GGAL — exclusivamente Puts, no Calls: el objetivo es protección de cola "
        "(convexidad ante caídas), no apuesta direccional alcista.",
    )
    _bullet(
        doc,
        f"Selección del strike: distancia porcentual fija al spot "
        f"(otm_pct = {moneyness['otm_pct']*100:.0f}%, parámetro placeholder — ver sección 2), "
        "no delta objetivo de Black-Scholes. Esta decisión es observable directamente en "
        "el dato de mercado y no depende de un modelo de pricing ni de la calidad de la "
        "volatilidad implícita.",
    )
    _bullet(
        doc,
        "Pérdida máxima por ciclo: la prima pagada (acotada y conocida de antemano). "
        "Ganancia potencial: no lineal, vía la convexidad de la opción (Gamma) si el "
        "activo cae fuerte antes del vencimiento.",
    )

    _heading(doc, "1.3 Mecánica de rebalanceo (event-driven)", 2)
    rb = config["rebalance"]
    _bullet(
        doc,
        "El polo agresivo compra un nuevo Put OTM apenas vence el contrato anterior — no "
        "hay calendario fijo de rebalanceo, el evento disparador es el vencimiento.",
    )
    _bullet(
        doc,
        f"Piso de días al vencimiento para entrar a un contrato: {rb['min_dias_vto']} días "
        "(evita comprar protección a punto de expirar, con poco valor extrínseco).",
    )
    _bullet(
        doc,
        f"Si no hay dato exacto en la fecha de entrada (ej. fin de semana), se busca hacia "
        f"adelante hasta {rb['max_gap_dias']} días.",
    )
    _bullet(
        doc,
        "Prioridad de fuente en el solapamiento sintético/real: "
        f"{' > '.join(rb['esquema_priority'])} — siempre se prefiere el dato observado "
        "sobre el reconstruido.",
    )

    _heading(doc, "1.4 Modelo de capitalización del polo agresivo", 2)
    _para(
        doc,
        "Surgió una pregunta metodológica concreta: si varios ciclos seguidos los Puts "
        "vencen sin valor (sin evento de cola), ¿se sigue invirtiendo un % fijo del "
        "capital corriente (erosión geométrica del principal) o solo los intereses "
        "generados por el T-Bill (principal intacto, posición más chica)? Se compararon "
        "3 reglas alternativas sobre el dataset real antes de fijar una para la tesis:",
    )
    cm = config["capital_models"]
    _bullet(
        doc,
        "interest_only: el presupuesto de cada ciclo es exactamente el interés devengado "
        "del T-Bill. El principal nunca se toca; en ciclos de tasa ~0 el presupuesto "
        "puede ser ~0 y no se compra protección ese ciclo.",
    )
    _bullet(
        doc,
        "fixed_ratio: el presupuesto es un % fijo (aggressive_weight) del capital total "
        "actual, recalculado cada ciclo. Si no hay eventos, el principal se erosiona "
        "geométricamente, pero el % invertido respecto al capital corriente es siempre "
        "el mismo.",
    )
    _bullet(
        doc,
        f"hybrid (modelo elegido para la tesis): el presupuesto usa el interés devengado "
        f"si alcanza un piso mínimo ({cm['min_position_pct']*100:.1f}% del capital total "
        "de ese ciclo); si no alcanza, completa con principal hasta ese piso. Combina lo "
        "mejor de los dos esquemas: preserva principal cuando el interés ya es suficiente, "
        "pero garantiza una posición mínima viable de protección incluso en ciclos de "
        "tasa baja.",
    )
    _para(
        doc,
        "En los 3 casos, el 'reset tras evento' es automático: si un ciclo paga fuerte, el "
        "capital total sube, y el próximo presupuesto (cualquiera sea la regla) se calcula "
        "sobre una base mayor — no hace falta una regla aparte para eso.",
    )

    # -------------------------------------------------------------------
    # 2. Supuestos metodológicos
    # -------------------------------------------------------------------
    _heading(doc, "2. Supuestos y decisiones metodológicas", 1)
    _para(
        doc,
        "Decisiones cerradas con el tutor académico (no deben modificarse sin "
        "discutirlas con él):",
    )
    _bullet(
        doc,
        "Subyacente: opción local de GGAL convertida a USD vía CCL (no derivado "
        "sintético sobre el ADR). Es data-driven, reproducible y defendible — evita el "
        "problema de pricear derivados teóricos sobre el ADR sin mercado observable.",
    )
    _bullet(
        doc,
        "Moneyness: % fijo de distancia entre spot y strike, no delta objetivo. Es "
        "observable directo, no depende de Black-Scholes y es robusto a errores en la "
        "volatilidad implícita.",
    )
    _bullet(
        doc,
        "Volatilidad: implícita por strike (no promedio), para preservar el skew/sonrisa "
        "— el promedio aplastaría la información del riesgo asimétrico, central para una "
        "tesis sobre cola izquierda.",
    )
    _bullet(
        doc,
        "Costos de transacción: modelados por liquidez. Volumen alto → spread bajo "
        f"({config['transaction_costs']['high_volume_spread']*100:.1f}%); volumen bajo → "
        f"spread alto ({config['transaction_costs']['low_volume_spread']*100:.0f}%), con "
        f"umbral de clasificación en {config['transaction_costs']['volume_threshold']:,.0f} "
        "ARS nominales operados.",
    )
    _bullet(
        doc,
        "Benchmarks: ambos, no uno solo. ADR GGAL buy-and-hold mide el efecto activo de "
        "la estrategia contra su propio subyacente; Merval en USD mide el efecto de "
        "mercado contra el equity argentino completo.",
    )

    _heading(doc, "2.1 Ventana temporal y reconstrucción sintética", 2)
    bt_cfg = config["backtest"]
    _para(
        doc,
        f"El backtest cubre {bt_cfg['start_date']} → {bt_cfg['end_date']}, dividido en dos "
        "tramos según la fuente de datos de opciones:",
    )
    _bullet(
        doc,
        "Tramo sintético (2019-01-01 → 2023-10-17): no existen datos reales de opciones "
        "de GGAL para este período. Se reconstruyen con Black-Scholes europeo, calibrando "
        "la superficie de skew (volatilidad implícita en función de moneyness y días al "
        "vencimiento) sobre los datos reales del tramo posterior y proyectándola hacia "
        "atrás. El spot histórico se toma del ADR de GGAL (NYSE) convertido a ARS vía CCL "
        "diario, y la tasa libre de riesgo de T-Bills 3M (FRED).",
    )
    _bullet(
        doc,
        "Tramo real (2023-10-18 → 2026-06-12): 17 archivos Historial de GGAL, formato "
        "tidy (una fila por contrato por fecha), con volatilidad implícita y griegas "
        "observadas directamente del mercado (o recalculadas con Black-Scholes calibrado "
        "contra el esquema que sí trae griegas, para los archivos más antiguos que no las "
        "incluyen).",
    )
    _para(
        doc,
        "Limitación reconocida: la calibración del skew se hace sobre 2023-2026 y se "
        "proyecta a 2019-2023, asumiendo que la forma de la sonrisa es estable aunque el "
        "nivel cambie. Es un supuesto fuerte — si en 2019 la crashophobia tenía una "
        "estructura distinta, los Puts OTM sintéticos de ese tramo están mal precieados. "
        "El recorte de 2015 a 2019 (en vez de arrancar en 2015) fue deliberado para "
        "limitar el horizonte sintético a ~5 años y mantener cobertura del evento PASO "
        "2019, el shock político más relevante del período.",
    )
    _para(
        doc,
        "Todos los resultados se reportan en USD: la conversión de pesos a dólares se "
        "hace vía el tipo de cambio Contado Con Liquidación (CCL), consistente para el "
        "subyacente, las opciones y el benchmark del Merval.",
    )

    # -------------------------------------------------------------------
    # 3. Métricas del backtest
    # -------------------------------------------------------------------
    _heading(doc, "3. Métricas analizadas en el backtest", 1)
    _para(
        doc,
        "Se calculan dos familias de métricas: de retorno/composición de capital (por "
        "ciclo de inversión, motor en `backtest.py`) y de riesgo ajustado (`metrics.py` / "
        "`benchmark.py`). La tesis combina ambas en una sola tabla comparativa.",
    )
    _heading(doc, "3.1 Métricas de retorno y composición", 2)
    _bullet(doc, "Capital final (USD): valor de la cartera al cierre de la ventana del backtest.")
    _bullet(doc, "Retorno total (%): variación del capital final respecto al capital inicial.")
    _bullet(
        doc,
        "CAGR (Compound Annual Growth Rate): retorno anualizado equivalente, para "
        "comparar estrategias con la misma ventana temporal independientemente de la "
        "frecuencia de los flujos.",
    )
    _bullet(
        doc,
        "Número de ciclos (n_cycles): cantidad de ciclos de Put completados durante el "
        "backtest.",
    )
    _bullet(
        doc,
        "Toques de principal (n_principal_touches): cantidad de ciclos en los que el "
        "presupuesto del polo agresivo excedió el interés devengado, es decir, se usó "
        "capital del principal (no solo intereses).",
    )
    _bullet(
        doc,
        "% de ciclos ITM (pct_cycles_itm): proporción de ciclos en los que el Put terminó "
        "in-the-money al vencimiento (el evento de cola se materializó y la posición "
        "pagó más que su costo de entrada).",
    )

    _heading(doc, "3.2 Métricas de riesgo ajustado", 2)
    _bullet(
        doc,
        "Max Drawdown (%): la mayor caída porcentual desde un máximo previo de la curva "
        "de equity — la métrica central de la tesis, porque mide exactamente lo que la "
        "Barbell busca evitar (riesgo de cola izquierda).",
    )
    _bullet(
        doc,
        "Calmar ratio: CAGR sobre el valor absoluto del Max Drawdown. A diferencia de "
        "Sharpe/Sortino (que penalizan la volatilidad de los retornos en general), el "
        "Calmar penaliza específicamente la peor caída histórica — el argumento de "
        "'retorno por unidad de riesgo de cola' que es el eje central de la tesis.",
    )
    _bullet(
        doc,
        "Sharpe ratio: retorno en exceso de la tasa libre de riesgo, sobre la "
        "volatilidad total de los retornos (penaliza tanto al alza como a la baja).",
    )
    _bullet(
        doc,
        "Sortino ratio: igual que Sharpe, pero solo penaliza la volatilidad a la baja "
        "(downside deviation) — más apropiado para una estrategia con retornos "
        "asimétricos por diseño, como la Barbell.",
    )
    _bullet(
        doc,
        "Expected Shortfall / CVaR (al 5%): pérdida promedio esperada en el peor 5% de "
        "los escenarios — complementa al Max Drawdown con una medida de cola que no "
        "depende de un único evento histórico.",
    )
    _para(
        doc,
        "Nota metodológica: la Barbell se anualiza usando la cadencia real de ciclos "
        f"(~{hybrid_full['periods_per_year']:.2f} ciclos/año, ciclos de ~2 meses), mientras "
        "que los benchmarks buy-and-hold se anualizan con 252 días de mercado por año. "
        "El Calmar ratio es comparable directamente entre ambos (usa CAGR y Max Drawdown, "
        "ya anuales); Sharpe y Sortino, en cambio, no son directamente comparables en "
        "magnitud entre la Barbell y los benchmarks por esta diferencia de frecuencia — "
        "se reportan a título informativo, no como comparación principal.",
    )

    # -------------------------------------------------------------------
    # 4. Resultados: Barbell vs. ADR vs. Merval
    # -------------------------------------------------------------------
    _heading(doc, "4. Resultados: Barbell (hybrid) vs. ADR GGAL vs. Merval USD", 1)
    _para(
        doc,
        f"Ventana del backtest: {bt_cfg['start_date']} a {bt_cfg['end_date']} "
        f"({hybrid_bt['n_cycles']:.0f} ciclos completos de Put para la Barbell). Capital "
        f"inicial: USD {cm['initial_capital']:,.0f}.",
    )

    display_cols = [
        "total_return_pct", "cagr", "max_drawdown_pct", "calmar", "sharpe", "sortino",
    ]
    display_df = full.loc[
        ["Barbell (hybrid)", "GGAL ADR buy-and-hold", "Merval en USD"], display_cols
    ].rename(columns={
        "total_return_pct": "Retorno total",
        "cagr": "CAGR",
        "max_drawdown_pct": "Max Drawdown",
        "calmar": "Calmar",
        "sharpe": "Sharpe",
        "sortino": "Sortino",
    })
    _table_from_df(doc, display_df, pct_cols=["Retorno total", "CAGR", "Max Drawdown"])

    doc.add_paragraph()
    _para(
        doc,
        f"Resultado central: la Barbell (hybrid) cae como máximo "
        f"{hybrid_full['max_drawdown_pct']:.1f}% desde su pico, frente a "
        f"{adr_full['max_drawdown_pct']:.1f}% del ADR GGAL y "
        f"{merval_full['max_drawdown_pct']:.1f}% del Merval en USD — una reducción de "
        "drawdown de un orden de magnitud. A cambio, la Barbell no supera el retorno "
        f"nominal del Merval ({merval_full['total_return_pct']:.1f}% vs. "
        f"{hybrid_full['total_return_pct']:.1f}%), aunque sí supera al ADR GGAL "
        f"({adr_full['total_return_pct']:.1f}%).",
        bold=True,
    )
    _para(
        doc,
        f"El Calmar ratio resume el trade-off: {hybrid_full['calmar']:.2f} para la "
        f"Barbell, frente a {adr_full['calmar']:.2f} del ADR y {merval_full['calmar']:.2f} "
        "del Merval — es decir, la Barbell genera entre 5 y 10 veces más retorno anual "
        "por cada punto de drawdown máximo soportado. Este es el argumento empírico "
        "central de la tesis: la convexidad de la estrategia funciona como un seguro "
        "contra la cola izquierda del equity argentino, no como un motor de retorno "
        "superior en términos nominales.",
    )

    _heading(doc, "4.1 Comparación con los otros dos modelos de capitalización", 2)
    _para(
        doc,
        "A título de referencia metodológica (no son parte de la comparación principal "
        "de la tesis, que usa el modelo hybrid), así se comportan las otras dos reglas de "
        "presupuesto sobre el mismo dataset:",
    )
    other_cols = ["total_return_pct", "cagr", "max_drawdown_pct", "calmar"]
    other_df = bt.loc[["interest_only", "fixed_ratio"], other_cols].rename(
        index={"interest_only": "interest_only (solo intereses)", "fixed_ratio": "fixed_ratio (% fijo del capital)"},
        columns={
            "total_return_pct": "Retorno total", "cagr": "CAGR",
            "max_drawdown_pct": "Max Drawdown", "calmar": "Calmar",
        },
    )
    _table_from_df(doc, other_df, pct_cols=["Retorno total", "CAGR", "Max Drawdown"])
    _para(
        doc,
        "fixed_ratio logra el mayor retorno nominal, pero a costa de un drawdown "
        f"({bt.loc['fixed_ratio', 'max_drawdown_pct']:.1f}%) comparable al de los "
        "benchmarks lineales — pierde la propiedad central de la Barbell de proteger el "
        "principal. interest_only protege el principal casi por completo (drawdown 0%), "
        "pero a costa de un retorno mucho menor en ciclos de tasa baja. hybrid es el punto "
        "medio elegido para la tesis: preserva la mayor parte de la protección de "
        "interest_only sin sacrificar tanto retorno.",
    )

    _heading(doc, "5. Figuras de referencia", 1)
    _para(
        doc,
        "Generadas por `src/report.py` en `reports/figures/`: `equity_curves.png` "
        "(curva de capital en escala lineal y logarítmica), `drawdown_curves.png` "
        "(drawdown desde el máximo previo de cada serie) y `summary_bars.png` (barras "
        "comparativas de retorno total, CAGR y max drawdown).",
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Memo guardado en: {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
